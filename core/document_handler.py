"""
core/document_handler.py — Safe Document Download & Analysis Module

Handles the interactive, safety-first workflow for documents discovered
during OSINT searches. NEVER auto-downloads — every document requires
explicit user consent. Downloads go to an isolated sandbox directory
and are parsed for text, metadata, and intelligence clues.

Document pipeline:
    1. Display document info to user
    2. Ask permission (B/X prompt)
    3. If approved: download to sandbox, validate, extract
    4. Scan extracted text for intelligence clues
    5. Return structured data for the Intelligence Engine
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path
from typing import Any, Optional

import httpx
from loguru import logger
from rich.console import Console
from rich.panel import Panel
from rich.table import Table
from rich.text import Text

from ui.locale_az import (
    DOC_FOUND_TITLE,
    DOC_FILENAME_LABEL,
    DOC_URL_LABEL,
    DOC_CONTEXT_LABEL,
    DOC_TYPE_LABEL,
    DOC_DOWNLOAD_PROMPT,
    DOC_DOWNLOADING,
    DOC_DOWNLOAD_SUCCESS,
    DOC_DOWNLOAD_ERROR,
    DOC_DOWNLOAD_SKIPPED,
    DOC_EXTRACTING,
    DOC_EXTRACT_SUCCESS,
    DOC_EXTRACT_ERROR,
    DOC_UNSUPPORTED_TYPE,
    DOC_SIZE_EXCEEDED,
    DOC_CLUES_FOUND,
    DOC_CLUE_EMAIL,
    DOC_CLUE_PHONE,
    DOC_CLUE_URL,
    DOC_CLUE_DATE,
    DOC_CLUE_NAME,
    DOC_METADATA_TITLE,
    DOC_METADATA_AUTHOR,
    DOC_METADATA_CREATED,
    DOC_METADATA_MODIFIED,
    DOC_METADATA_CREATOR,
    DOC_CLEANUP_DONE,
    DOC_SANDBOX_CREATED,
)

# ---------------------------------------------------------------------------
# Supported document types and size limits
# ---------------------------------------------------------------------------
_SUPPORTED_TYPES: dict[str, str] = {
    ".pdf": "PDF",
    ".docx": "DOCX",
    ".doc": "DOC (legacy)",
    ".xlsx": "XLSX",
    ".xls": "XLS (legacy)",
    ".csv": "CSV",
    ".txt": "Text",
    ".rtf": "RTF",
}

_DEFAULT_MAX_FILE_SIZE_MB: int = 50


class DocumentHandler:
    """Safe document download and analysis handler.

    Implements the interactive document approval workflow: displays document
    info to the user, asks for explicit consent before downloading, downloads
    to an isolated sandbox, extracts text and metadata, and scans for
    intelligence clues.

    Args:
        config: Document handling configuration from settings.yaml.
            Expected keys: sandbox_dir, allowed_types, auto_download,
            max_file_size_mb, clean_on_exit.
        console: Rich Console instance for interactive output.
    """

    def __init__(self, config: dict, console: Console) -> None:
        self._config = config
        self._console = console

        # Configuration
        self._sandbox_dir = Path(config.get("sandbox_dir", "./sandbox"))
        self._allowed_types: list[str] = config.get(
            "allowed_types", ["pdf", "docx", "xlsx"]
        )
        self._auto_download: bool = config.get("auto_download", False)  # MUST remain False
        self._max_file_size_mb: int = config.get("max_file_size_mb", _DEFAULT_MAX_FILE_SIZE_MB)
        self._clean_on_exit: bool = config.get("clean_on_exit", True)

        # Tracking
        self._downloaded_files: list[Path] = []
        self._skipped_documents: list[dict[str, str]] = []
        self._processed_count: int = 0

        # Ensure sandbox exists
        self._ensure_sandbox()

        logger.debug(
            "DocumentHandler initialized: sandbox={}, allowed_types={}, max_size={}MB",
            self._sandbox_dir,
            self._allowed_types,
            self._max_file_size_mb,
        )

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    async def process_document(self, doc_info: dict[str, str]) -> Optional[dict[str, Any]]:
        """Show document to user, ask permission, download and extract if approved.

        This is the main entry point for the interactive document workflow.
        NEVER downloads automatically — always asks the user first.

        Args:
            doc_info: Dictionary with document information.
                Expected keys: url, title, snippet, document_type.

        Returns:
            Extracted data dict if processed, None if skipped/failed.
            On success: {text, metadata, clues, filepath, url, title}.
        """
        self._processed_count += 1

        url = doc_info.get("url", "")
        title = doc_info.get("title", "N/A")
        snippet = doc_info.get("snippet", "")
        doc_type = doc_info.get("document_type", "unknown")

        # --- Step 1: Display document info panel ---
        self._display_document_info(url, title, snippet, doc_type)

        # --- Step 2: Ask user for permission ---
        if self._auto_download:
            # Safety override — auto_download should NEVER be True in production
            logger.critical(
                "auto_download is True! This violates safety rules. Forcing manual approval."
            )

        try:
            user_input = self._console.input(
                f"\n  {DOC_DOWNLOAD_PROMPT} "
            ).strip().upper()
        except (EOFError, KeyboardInterrupt):
            user_input = "X"

        # --- Step 3: Handle user decision ---
        if user_input not in ("B", "Y", "YES", "BƏLİ", "BELI"):
            # User declined — log and move on
            self._skipped_documents.append({"url": url, "title": title, "type": doc_type})
            self._console.print(f"  [dim]{DOC_DOWNLOAD_SKIPPED}[/dim]")
            logger.info("Document skipped by user: {} ({})", url[:80], doc_type)
            return None

        # --- Step 4: Download to sandbox ---
        filename = self._generate_safe_filename(url, doc_type)
        filepath = await self.download_to_sandbox(url, filename)

        if not filepath:
            return None

        # --- Step 5: Validate file type and size ---
        file_ext = filepath.suffix.lower()
        if file_ext.lstrip(".") not in self._allowed_types:
            self._console.print(
                f"  [yellow]⚠[/yellow] {DOC_UNSUPPORTED_TYPE.format(file_type=file_ext)}"
            )
            logger.warning("Unsupported file type: {}", file_ext)
            filepath.unlink(missing_ok=True)
            return None

        # --- Step 6: Extract text and metadata ---
        self._console.print(f"  [cyan]⟳[/cyan] {DOC_EXTRACTING}")
        extracted = self._extract_document(filepath)

        if not extracted:
            return None

        # --- Step 7: Scan for intelligence clues ---
        text = extracted.get("text", "")
        clues = self.scan_for_clues(text)

        if clues:
            self._display_clues(clues)

        # --- Build result ---
        result: dict[str, Any] = {
            "text": text,
            "metadata": extracted.get("metadata", {}),
            "clues": clues,
            "filepath": str(filepath),
            "url": url,
            "title": title,
            "document_type": doc_type,
        }

        self._console.print(f"  [green]✓[/green] {DOC_EXTRACT_SUCCESS}")
        logger.info(
            "Document processed: {} | {} chars text | {} clues found",
            filepath.name,
            len(text),
            len(clues),
        )

        return result

    async def download_to_sandbox(
        self,
        url: str,
        filename: str,
    ) -> Optional[Path]:
        """Download a file to the sandbox directory safely.

        Uses httpx with streaming to monitor file size during download.
        Validates file type and size after download.

        Args:
            url: URL to download from.
            filename: Filename to save as in the sandbox.

        Returns:
            Path to the downloaded file, or None on failure.
        """
        self._ensure_sandbox()
        filepath = self._sandbox_dir / filename

        self._console.print(f"  [cyan]⟳[/cyan] {DOC_DOWNLOADING}")
        logger.info("Downloading document: {} -> {}", url[:80], filepath)

        try:
            async with httpx.AsyncClient(
                timeout=httpx.Timeout(60.0),
                follow_redirects=True,
                headers={
                    "Accept": "*/*",
                    "Accept-Language": "en-US,en;q=0.9",
                },
            ) as client:
                async with client.stream("GET", url) as response:
                    response.raise_for_status()

                    # Check Content-Length if available
                    content_length = response.headers.get("content-length")
                    if content_length:
                        size_mb = int(content_length) / (1024 * 1024)
                        if size_mb > self._max_file_size_mb:
                            self._console.print(
                                f"  [red]✗[/red] {DOC_SIZE_EXCEEDED.format(size=size_mb, max_size=self._max_file_size_mb)}"
                            )
                            logger.warning(
                                "File too large: {:.1f}MB (max {}MB)",
                                size_mb,
                                self._max_file_size_mb,
                            )
                            return None

                    # Stream download with size monitoring
                    downloaded_bytes = 0
                    max_bytes = self._max_file_size_mb * 1024 * 1024

                    with open(filepath, "wb") as f:
                        async for chunk in response.aiter_bytes(8192):
                            downloaded_bytes += len(chunk)
                            if downloaded_bytes > max_bytes:
                                self._console.print(
                                    f"  [red]✗[/red] {DOC_SIZE_EXCEEDED.format(size=downloaded_bytes / (1024*1024), max_size=self._max_file_size_mb)}"
                                )
                                filepath.unlink(missing_ok=True)
                                return None
                            f.write(chunk)

            self._downloaded_files.append(filepath)
            size_kb = filepath.stat().st_size / 1024
            self._console.print(
                f"  [green]✓[/green] {DOC_DOWNLOAD_SUCCESS.format(filename=filename, size=size_kb)}"
            )
            logger.info("Downloaded: {} ({:.1f} KB)", filename, size_kb)
            return filepath

        except httpx.HTTPStatusError as exc:
            logger.error("HTTP error downloading {}: {}", url[:80], exc.response.status_code)
            self._console.print(f"  [red]✗[/red] {DOC_DOWNLOAD_ERROR}: HTTP {exc.response.status_code}")
            return None
        except httpx.TimeoutException:
            logger.error("Timeout downloading: {}", url[:80])
            self._console.print(f"  [red]✗[/red] {DOC_DOWNLOAD_ERROR}: Timeout")
            return None
        except Exception as exc:
            logger.error("Download error: {}", exc)
            self._console.print(f"  [red]✗[/red] {DOC_DOWNLOAD_ERROR}: {exc}")
            return None

    def extract_pdf(self, filepath: Path) -> dict[str, Any]:
        """Extract text and metadata from a PDF using pdfplumber.

        Safe extraction — no code execution, only text and metadata.

        Args:
            filepath: Path to the PDF file.

        Returns:
            Dict with keys: text (str), metadata (dict).
        """
        try:
            import pdfplumber

            text_parts: list[str] = []
            metadata: dict[str, Any] = {}

            with pdfplumber.open(filepath) as pdf:
                # Extract metadata
                if pdf.metadata:
                    metadata = {
                        "author": pdf.metadata.get("Author", ""),
                        "creator": pdf.metadata.get("Creator", ""),
                        "producer": pdf.metadata.get("Producer", ""),
                        "created": pdf.metadata.get("CreationDate", ""),
                        "modified": pdf.metadata.get("ModDate", ""),
                        "title": pdf.metadata.get("Title", ""),
                        "subject": pdf.metadata.get("Subject", ""),
                        "pages": len(pdf.pages),
                    }

                # Extract text from all pages
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            logger.info("PDF extracted: {} pages, {} chars", metadata.get("pages", 0), sum(len(t) for t in text_parts))
            return {"text": "\n\n".join(text_parts), "metadata": metadata}

        except ImportError:
            logger.error("pdfplumber not installed. Install with: pip install pdfplumber")
            return {"text": "", "metadata": {"error": "pdfplumber not installed"}}
        except Exception as exc:
            logger.error("PDF extraction error: {}", exc)
            return {"text": "", "metadata": {"error": str(exc)}}

    def extract_docx(self, filepath: Path) -> dict[str, Any]:
        """Extract text and metadata from a DOCX using python-docx.

        Args:
            filepath: Path to the DOCX file.

        Returns:
            Dict with keys: text (str), metadata (dict).
        """
        try:
            from docx import Document

            doc = Document(filepath)

            # Extract metadata from core properties
            metadata: dict[str, Any] = {}
            if doc.core_properties:
                props = doc.core_properties
                metadata = {
                    "author": props.author or "",
                    "created": str(props.created) if props.created else "",
                    "modified": str(props.modified) if props.modified else "",
                    "last_modified_by": props.last_modified_by or "",
                    "title": props.title or "",
                    "subject": props.subject or "",
                    "category": props.category or "",
                    "revision": props.revision,
                }

            # Extract text from paragraphs
            text_parts: list[str] = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_texts: list[str] = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_texts.append(cell.text.strip())
                    if row_texts:
                        text_parts.append(" | ".join(row_texts))

            logger.info("DOCX extracted: {} paragraphs, {} chars", len(text_parts), sum(len(t) for t in text_parts))
            return {"text": "\n".join(text_parts), "metadata": metadata}

        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return {"text": "", "metadata": {"error": "python-docx not installed"}}
        except Exception as exc:
            logger.error("DOCX extraction error: {}", exc)
            return {"text": "", "metadata": {"error": str(exc)}}

    def extract_xlsx(self, filepath: Path) -> dict[str, Any]:
        """Extract data and metadata from an XLSX using openpyxl.

        Args:
            filepath: Path to the XLSX file.

        Returns:
            Dict with keys: text (str), metadata (dict).
        """
        try:
            from openpyxl import load_workbook

            wb = load_workbook(filepath, read_only=True, data_only=True)

            # Extract metadata
            metadata: dict[str, Any] = {}
            if wb.properties:
                props = wb.properties
                metadata = {
                    "author": props.creator or "",
                    "created": str(props.created) if props.created else "",
                    "modified": str(props.modified) if props.modified else "",
                    "last_modified_by": props.lastModifiedBy or "",
                    "title": props.title or "",
                    "subject": props.subject or "",
                    "sheets": wb.sheetnames,
                }

            # Extract text from all sheets
            text_parts: list[str] = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                text_parts.append(f"--- {sheet_name} ---")
                for row in ws.iter_rows(values_only=True):
                    cell_values = [str(v) for v in row if v is not None]
                    if cell_values:
                        text_parts.append(" | ".join(cell_values))

            wb.close()

            logger.info("XLSX extracted: {} sheets, {} rows", len(wb.sheetnames), len(text_parts))
            return {"text": "\n".join(text_parts), "metadata": metadata}

        except ImportError:
            logger.error("openpyxl not installed. Install with: pip install openpyxl")
            return {"text": "", "metadata": {"error": "openpyxl not installed"}}
        except Exception as exc:
            logger.error("XLSX extraction error: {}", exc)
            return {"text": "", "metadata": {"error": str(exc)}}

    def scan_for_clues(self, text: str) -> list[dict[str, str]]:
        """Scan extracted text for intelligence clues.

        Uses regex patterns to find potentially valuable data points:
        email addresses, phone numbers (especially +994 AZ format),
        URLs, dates, and potential names.

        Args:
            text: The extracted text to scan.

        Returns:
            List of clue dicts: {type, value, context}.
            Context is a short surrounding text snippet for verification.
        """
        if not text:
            return []

        clues: list[dict[str, str]] = []
        seen_values: set[str] = set()  # Deduplication

        # --- Email addresses ---
        email_pattern = re.compile(
            r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}',
            re.IGNORECASE,
        )
        for match in email_pattern.finditer(text):
            email = match.group().lower()
            if email not in seen_values:
                seen_values.add(email)
                context = self._extract_context(text, match.start(), match.end())
                clues.append({
                    "type": "email",
                    "value": email,
                    "context": context,
                    "label": DOC_CLUE_EMAIL,
                })

        # --- Phone numbers (with special handling for Azerbaijan +994) ---
        phone_patterns = [
            # Azerbaijan format: +994 XX XXX XX XX
            re.compile(r'\+994[\s\-]?\d{2}[\s\-]?\d{3}[\s\-]?\d{2}[\s\-]?\d{2}'),
            # International format
            re.compile(r'\+\d{1,3}[\s\-]?\(?\d{2,4}\)?[\s\-]?\d{3,4}[\s\-]?\d{2,4}[\s\-]?\d{0,4}'),
            # Local formats: (0XX) XXX-XX-XX
            re.compile(r'\(0\d{2}\)[\s\-]?\d{3}[\s\-]?\d{2}[\s\-]?\d{2}'),
        ]
        for pattern in phone_patterns:
            for match in pattern.finditer(text):
                phone = match.group().strip()
                # Normalize by removing spaces/dashes for dedup
                normalized = re.sub(r'[\s\-\(\)]', '', phone)
                if normalized not in seen_values and len(normalized) >= 7:
                    seen_values.add(normalized)
                    context = self._extract_context(text, match.start(), match.end())
                    clues.append({
                        "type": "phone",
                        "value": phone,
                        "context": context,
                        "label": DOC_CLUE_PHONE,
                    })

        # --- URLs ---
        url_pattern = re.compile(
            r'https?://[^\s<>"\')\]]+',
            re.IGNORECASE,
        )
        for match in url_pattern.finditer(text):
            url = match.group().rstrip(".,;:)")
            if url not in seen_values and len(url) > 10:
                seen_values.add(url)
                context = self._extract_context(text, match.start(), match.end())
                clues.append({
                    "type": "url",
                    "value": url,
                    "context": context,
                    "label": DOC_CLUE_URL,
                })

        # --- Dates and years ---
        date_patterns = [
            # DD.MM.YYYY or DD/MM/YYYY
            re.compile(r'\b\d{1,2}[./\-]\d{1,2}[./\-]\d{4}\b'),
            # YYYY-MM-DD (ISO)
            re.compile(r'\b\d{4}[.\-/]\d{1,2}[.\-/]\d{1,2}\b'),
            # Standalone years (1980-2030)
            re.compile(r'\b(19[89]\d|20[012]\d)\b'),
        ]
        for pattern in date_patterns:
            for match in pattern.finditer(text):
                date_val = match.group()
                if date_val not in seen_values:
                    seen_values.add(date_val)
                    context = self._extract_context(text, match.start(), match.end())
                    clues.append({
                        "type": "date",
                        "value": date_val,
                        "context": context,
                        "label": DOC_CLUE_DATE,
                    })

        # --- Potential names (capitalized word sequences, 2-3 words) ---
        name_pattern = re.compile(
            r'\b([A-ZÇƏĞIİÖŞÜА-Я][a-zçəğıiöşüа-яё]+(?:\s+[A-ZÇƏĞIİÖŞÜА-Я][a-zçəğıiöşüа-яё]+){1,2})\b'
        )
        for match in name_pattern.finditer(text):
            name = match.group().strip()
            if name not in seen_values and len(name) > 4:
                # Filter out common non-name phrases
                lower_name = name.lower()
                skip_words = {
                    "the", "and", "for", "this", "that", "with", "from",
                    "january", "february", "march", "april", "may", "june",
                    "july", "august", "september", "october", "november", "december",
                }
                words = lower_name.split()
                if not any(w in skip_words for w in words):
                    seen_values.add(name)
                    context = self._extract_context(text, match.start(), match.end())
                    clues.append({
                        "type": "name",
                        "value": name,
                        "context": context,
                        "label": DOC_CLUE_NAME,
                    })

        logger.info("Scanned text for clues: {} found", len(clues))
        return clues

    def cleanup(self) -> None:
        """Clean the sandbox directory, removing all downloaded files.

        Called on session exit if clean_on_exit is enabled in config.
        """
        if not self._clean_on_exit:
            logger.info("Sandbox cleanup skipped (clean_on_exit=False).")
            return

        try:
            if self._sandbox_dir.exists():
                shutil.rmtree(self._sandbox_dir)
                self._console.print(f"  [dim]{DOC_CLEANUP_DONE}[/dim]")
                logger.info("Sandbox cleaned: {}", self._sandbox_dir)
        except Exception as exc:
            logger.error("Sandbox cleanup error: {}", exc)

        self._downloaded_files.clear()

    @property
    def downloaded_count(self) -> int:
        """Number of documents downloaded in this session."""
        return len(self._downloaded_files)

    @property
    def skipped_count(self) -> int:
        """Number of documents skipped by user in this session."""
        return len(self._skipped_documents)

    @property
    def skipped_documents(self) -> list[dict[str, str]]:
        """List of documents the user declined to download."""
        return list(self._skipped_documents)

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _extract_document(self, filepath: Path) -> Optional[dict[str, Any]]:
        """Route document extraction to the appropriate handler based on file type.

        Args:
            filepath: Path to the downloaded document.

        Returns:
            Extracted data dict, or None if extraction failed.
        """
        ext = filepath.suffix.lower()

        try:
            if ext == ".pdf":
                return self.extract_pdf(filepath)
            elif ext in (".docx", ".doc"):
                return self.extract_docx(filepath)
            elif ext in (".xlsx", ".xls"):
                return self.extract_xlsx(filepath)
            elif ext in (".csv", ".txt"):
                return self._extract_text(filepath)
            else:
                self._console.print(
                    f"  [yellow]⚠[/yellow] {DOC_UNSUPPORTED_TYPE.format(file_type=ext)}"
                )
                return None
        except Exception as exc:
            logger.error("Document extraction error for {}: {}", filepath, exc)
            self._console.print(f"  [red]✗[/red] {DOC_EXTRACT_ERROR}: {exc}")
            return None

    def _extract_text(self, filepath: Path) -> dict[str, Any]:
        """Extract text from a plain text or CSV file.

        Args:
            filepath: Path to the text file.

        Returns:
            Dict with keys: text, metadata.
        """
        try:
            # Try UTF-8 first, fall back to latin-1
            for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
                try:
                    text = filepath.read_text(encoding=encoding)
                    return {
                        "text": text,
                        "metadata": {
                            "encoding": encoding,
                            "size_bytes": filepath.stat().st_size,
                        },
                    }
                except UnicodeDecodeError:
                    continue

            return {"text": "", "metadata": {"error": "Unable to decode file"}}

        except Exception as exc:
            logger.error("Text extraction error: {}", exc)
            return {"text": "", "metadata": {"error": str(exc)}}

    def _ensure_sandbox(self) -> None:
        """Create the sandbox directory if it doesn't exist."""
        if not self._sandbox_dir.exists():
            self._sandbox_dir.mkdir(parents=True, exist_ok=True)
            logger.info(DOC_SANDBOX_CREATED, self._sandbox_dir)

    def _generate_safe_filename(self, url: str, doc_type: str) -> str:
        """Generate a safe filename from a URL and document type.

        Strips unsafe characters and ensures a valid filename.

        Args:
            url: Source URL of the document.
            doc_type: Detected document type (pdf, docx, etc.).

        Returns:
            Safe filename string.
        """
        from urllib.parse import urlparse, unquote
        import hashlib

        parsed = urlparse(url)
        path = unquote(parsed.path)

        # Try to extract original filename from URL path
        if "/" in path:
            original_name = path.rsplit("/", 1)[-1]
        else:
            original_name = path

        # Clean the filename
        safe_name = re.sub(r'[^\w\-.]', '_', original_name)

        if not safe_name or len(safe_name) < 3:
            # Generate from URL hash
            url_hash = hashlib.md5(url.encode()).hexdigest()[:8]
            ext = f".{doc_type}" if doc_type and not doc_type.startswith(".") else (doc_type or ".bin")
            safe_name = f"doc_{url_hash}{ext}"

        # Ensure proper extension
        if not any(safe_name.lower().endswith(ext) for ext in _SUPPORTED_TYPES):
            if doc_type:
                ext = doc_type if doc_type.startswith(".") else f".{doc_type}"
                safe_name += ext

        return safe_name

    def _display_document_info(
        self,
        url: str,
        title: str,
        snippet: str,
        doc_type: str,
    ) -> None:
        """Display document information in a formatted panel.

        Args:
            url: Document URL.
            title: Document title from search results.
            snippet: Context snippet from search results.
            doc_type: Detected document type.
        """
        table = Table(show_header=False, padding=(0, 2), box=None, expand=False)
        table.add_column("Label", style="bright_cyan", width=16)
        table.add_column("Value", style="white")

        table.add_row(DOC_FILENAME_LABEL, title or "N/A")
        table.add_row(DOC_TYPE_LABEL, doc_type.upper() if doc_type else "N/A")
        table.add_row(DOC_URL_LABEL, f"[link={url}]{url[:80]}...[/link]" if len(url) > 80 else f"[link={url}]{url}[/link]")
        if snippet:
            table.add_row(DOC_CONTEXT_LABEL, snippet[:120])

        self._console.print()
        self._console.print(
            Panel(
                table,
                title=f"📄 {DOC_FOUND_TITLE}",
                border_style="yellow",
                padding=(1, 2),
            )
        )

    def _display_clues(self, clues: list[dict[str, str]]) -> None:
        """Display discovered intelligence clues in a table.

        Args:
            clues: List of clue dicts from scan_for_clues.
        """
        self._console.print(f"\n  [yellow]🔎[/yellow] {DOC_CLUES_FOUND.format(count=len(clues))}")

        table = Table(show_header=True, header_style="bold yellow")
        table.add_column("Növ", style="cyan", width=10)
        table.add_column("Dəyər", style="bright_white")
        table.add_column("Kontekst", style="dim", width=40)

        for clue in clues[:15]:  # Show top 15 clues
            table.add_row(
                clue.get("label", clue["type"]),
                clue["value"],
                clue.get("context", "")[:40],
            )

        if len(clues) > 15:
            table.add_row("...", f"+{len(clues) - 15} daha", "", style="dim")

        self._console.print(table)

    @staticmethod
    def _extract_context(text: str, start: int, end: int, window: int = 40) -> str:
        """Extract surrounding context around a match in the text.

        Args:
            text: Full text.
            start: Match start index.
            end: Match end index.
            window: Number of characters to include before/after.

        Returns:
            Context snippet string.
        """
        ctx_start = max(0, start - window)
        ctx_end = min(len(text), end + window)
        context = text[ctx_start:ctx_end].strip()
        # Clean up whitespace
        context = re.sub(r'\s+', ' ', context)
        if ctx_start > 0:
            context = "..." + context
        if ctx_end < len(text):
            context = context + "..."
        return context

    def _display_metadata(self, metadata: dict[str, Any]) -> None:
        """Display document metadata in a formatted panel.

        Args:
            metadata: Extracted metadata dictionary.
        """
        if not metadata or metadata.get("error"):
            return

        table = Table(show_header=False, padding=(0, 2), box=None)
        table.add_column("Field", style="bright_cyan", width=18)
        table.add_column("Value", style="white")

        field_map = {
            "author": DOC_METADATA_AUTHOR,
            "created": DOC_METADATA_CREATED,
            "modified": DOC_METADATA_MODIFIED,
            "creator": DOC_METADATA_CREATOR,
        }

        for key, label in field_map.items():
            value = metadata.get(key, "")
            if value:
                table.add_row(label, str(value))

        if any(metadata.get(k) for k in field_map):
            self._console.print(
                Panel(
                    table,
                    title=DOC_METADATA_TITLE,
                    border_style="cyan",
                    padding=(0, 2),
                )
            )
